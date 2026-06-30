from __future__ import annotations

import json
from pathlib import Path

import joblib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sklearn.model_selection import train_test_split


BASE_CSV = Path("Base_ScoreCredito_QuantumFinance(8).csv")
RESULTS_JSON = Path("outputs/resultados_modelo_credit_scoring.json")
MODEL_PATH = Path("outputs/modelo_credit_scoring_linear.joblib")
OUT_DOCX = Path("outputs/Relatorio_QuantumFinance_CreditScoring.docx")
IMG_DIR = Path("outputs/figuras")


def load_inputs():
    df = pd.read_csv(BASE_CSV, sep=";", decimal=",")
    results = json.loads(RESULTS_JSON.read_text(encoding="utf-8"))
    return df, results


def format_num(value, digits=2):
    if isinstance(value, str):
        return value
    return f"{value:,.{digits}f}".replace(",", "X").replace(".", ",").replace("X", ".")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, "F2F4F7")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            if widths:
                set_cell_width(cell, widths[col_idx])


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, widths)
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def save_charts(df, results):
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    plt.style.use("seaborn-v0_8-whitegrid")

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist(df["SCORE_CREDITO"], bins=35, color="#2E74B5", edgecolor="white")
    ax.set_title("Distribuicao do score de credito")
    ax.set_xlabel("SCORE_CREDITO")
    ax.set_ylabel("Frequencia")
    fig.tight_layout()
    fig.savefig(IMG_DIR / "hist_score.png", dpi=180)
    plt.close(fig)

    corr = pd.Series(results["correlacao_score"]).drop("SCORE_CREDITO").sort_values()
    fig, ax = plt.subplots(figsize=(7, 4.6))
    colors = ["#9B1C1C" if v < 0 else "#2E74B5" for v in corr.values]
    ax.barh(corr.index, corr.values, color=colors)
    ax.axvline(0, color="#555555", linewidth=0.8)
    ax.set_title("Correlacao com SCORE_CREDITO")
    ax.set_xlabel("Correlacao de Pearson")
    fig.tight_layout()
    fig.savefig(IMG_DIR / "corr_score.png", dpi=180)
    plt.close(fig)

    model = joblib.load(MODEL_PATH)
    X = df.drop(columns=["SCORE_CREDITO", "id"])
    y = df["SCORE_CREDITO"]
    _, X_test, _, y_test = train_test_split(X, y, test_size=0.25, random_state=42)
    pred = model.predict(X_test)
    fig, ax = plt.subplots(figsize=(5.8, 5))
    ax.scatter(y_test, pred, alpha=0.35, color="#1F4D78", s=14)
    lims = [min(y_test.min(), pred.min()), max(y_test.max(), pred.max())]
    ax.plot(lims, lims, color="#9B1C1C", linewidth=1.2)
    ax.set_title("Regressao linear: real vs. previsto")
    ax.set_xlabel("Score real")
    ax.set_ylabel("Score previsto")
    fig.tight_layout()
    fig.savefig(IMG_DIR / "real_vs_previsto.png", dpi=180)
    plt.close(fig)


def build_doc():
    df, results = load_inputs()
    save_charts(df, results)

    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Relatorio Estatistico - Credit Scoring Quantum Finance")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.add_run("Base analisada: Base_ScoreCredito_QuantumFinance.csv | Amostra: ")
    subtitle.add_run(f"{results['n_linhas']:,}".replace(",", ".")).bold = True
    subtitle.add_run(" clientes | Alvo: SCORE_CREDITO")

    doc.add_heading("Resumo executivo", level=1)
    add_bullet(doc, "A regressao linear multipla foi o modelo mais adequado entre os dois modelos testados, com R2 de 0,626 no teste e R2 medio de 0,656 em validacao cruzada.")
    add_bullet(doc, "A arvore de decisao teve desempenho inferior, com R2 de 0,564 e erro maior; ela e util como benchmark interpretavel, mas nao superou a regressao neste conjunto.")
    add_bullet(doc, "As variaveis com maior peso preditivo foram valor do imovel, casa propria, tempo no ultimo servico, salario, quantidade de carros e quantidade de cartoes.")
    add_bullet(doc, "As suposicoes da regressao foram atendidas apenas parcialmente: a independencia dos residuos e adequada, mas normalidade e homocedasticidade foram rejeitadas pelos testes formais.")

    doc.add_heading("Quadro conceitual estatistico", level=1)
    rows = [
        ("Tema", "Credit scoring para credito ao consumidor."),
        ("Problema", "A Quantum Finance precisa aprimorar a concessao de credito diante do aumento de inadimplencia e quer identificar clientes de menor risco."),
        ("Hipoteses conceituais", "Clientes com maior patrimonio, renda, estabilidade profissional e ativos tendem a apresentar maior SCORE_CREDITO; regioes de moradia e atributos cadastrais tambem podem influenciar o risco."),
        ("Objetivo principal", "Construir e avaliar um modelo preditivo para estimar o SCORE_CREDITO e apoiar analistas de credito e gerentes de conta."),
        ("Populacao de estudo", "Carteira atual de clientes da Quantum Finance, com 10.127 registros e 15 colunas no arquivo disponibilizado."),
        ("Plano basico de analise", "Tratamento da base; analise descritiva; correlacao; treino/teste; regressao linear multipla; arvore de decisao; avaliacao por MAE, RMSE, MAPE, R2 e validacao cruzada."),
        ("Tecnica estatistica", "Regressao linear multipla com variaveis numericas padronizadas e variaveis categoricas codificadas por one-hot encoding; comparacao com arvore de decisao regressora."),
        ("Resultado principal", "A regressao linear apresentou melhor equilibrio entre erro, explicabilidade e estabilidade, explicando cerca de 62,6% da variacao do score no teste."),
    ]
    add_table(doc, ["Componente", "Descricao"], rows, widths=[2300, 7060])

    doc.add_heading("Analise descritiva das variaveis", level=1)
    doc.add_paragraph("A base nao possui valores nulos. O campo estado_civil possui uma categoria 'na', que foi mantida como categoria informativa no modelo.")
    desc = pd.DataFrame(results["descritiva_numerica"]).T
    selected = ["idade", "Qte_dependentes", "tempo_ultimoservico", "vl_salario_mil", "vl_imovel_em_mil", "Qte_cartoes", "Qte_carros", "SCORE_CREDITO"]
    rows = []
    for var in selected:
        row = desc.loc[var]
        rows.append([var, format_num(row["mean"]), format_num(row["std"]), format_num(row["min"]), format_num(row["50%"]), format_num(row["max"])])
    add_table(doc, ["Variavel", "Media", "Desvio", "Min", "Mediana", "Max"], rows, widths=[2100, 1450, 1450, 1450, 1450, 1460])
    doc.add_picture(str(IMG_DIR / "hist_score.png"), width=Inches(6.2))

    doc.add_heading("Distribuicao das variaveis categoricas", level=2)
    cat_rows = []
    for col, counts in results["categorias"].items():
        total = sum(counts.values())
        top = sorted(counts.items(), key=lambda kv: kv[1], reverse=True)
        resumo = "; ".join([f"{k}: {v} ({v / total:.1%})".replace(".", ",") for k, v in top])
        cat_rows.append([col, resumo])
    add_table(doc, ["Variavel", "Distribuicao"], cat_rows, widths=[2100, 7260])

    doc.add_heading("Analise de correlacao", level=1)
    corr = pd.Series(results["correlacao_score"]).drop("SCORE_CREDITO").sort_values(ascending=False)
    rows = [[idx, format_num(val, 3)] for idx, val in corr.items()]
    add_table(doc, ["Variavel numerica", "Correlacao com SCORE_CREDITO"], rows, widths=[5200, 4160])
    doc.add_paragraph("As maiores correlacoes positivas com o score aparecem em valor do imovel, salario, tempo no ultimo servico, quantidade de cartoes, quantidade de carros e casa propria. A variavel reg_moradia tem correlacao negativa moderada.")
    doc.add_picture(str(IMG_DIR / "corr_score.png"), width=Inches(6.2))

    doc.add_heading("Modelo adequado para resolver o problema", level=1)
    doc.add_paragraph("O modelo recomendado para esta etapa e a regressao linear multipla. A escolha se justifica por tres motivos: melhor metrica no teste, melhor R2 medio em validacao cruzada e maior interpretabilidade para decisao de credito.")
    metrics_rows = []
    for nome, metricas in results["modelos"].items():
        metrics_rows.append([
            nome,
            format_num(metricas["MAE"]),
            format_num(metricas["RMSE"]),
            format_num(metricas["R2"], 3),
            f"{metricas['MAPE']:.2%}".replace(".", ","),
            format_num(metricas["CV_R2_media"], 3),
        ])
    add_table(doc, ["Modelo", "MAE", "RMSE", "R2 teste", "MAPE", "R2 CV medio"], metrics_rows, widths=[2600, 1300, 1300, 1300, 1300, 1560])
    doc.add_picture(str(IMG_DIR / "real_vs_previsto.png"), width=Inches(5.4))

    doc.add_heading("Resultados e qualidade do modelo", level=1)
    lin = results["modelos"]["Regressao Linear"]
    arv = results["modelos"]["Arvore de Decisao"]
    doc.add_paragraph(
        f"A regressao linear obteve MAE de {format_num(lin['MAE'])} pontos de score e RMSE de {format_num(lin['RMSE'])}. "
        f"Na pratica, o erro medio absoluto indica que a previsao individual costuma se afastar cerca de {format_num(lin['MAE'])} pontos do score observado."
    )
    doc.add_paragraph(
        f"A arvore apresentou MAE de {format_num(arv['MAE'])}, RMSE de {format_num(arv['RMSE'])} e R2 de {format_num(arv['R2'], 3)}, ficando abaixo da regressao. "
        "A comparacao sugere que, com estes atributos, a relacao linear captura melhor o padrao geral da base do que uma arvore rasa."
    )

    doc.add_heading("Diagnostico das suposicoes da regressao", level=2)
    diag = results["diagnostico_regressao"]
    diag_rows = [
        ["Independencia dos residuos", f"Durbin-Watson = {format_num(diag['Durbin_Watson'], 3)}", "Adequada; valor proximo de 2 indica baixa autocorrelacao."],
        ["Normalidade dos residuos", f"Jarque-Bera p-valor = {diag['Jarque_Bera_pvalor']:.2e}", "Nao atendida pelo teste formal."],
        ["Homoscedasticidade", f"Breusch-Pagan p-valor = {diag['Breusch_Pagan_pvalor']:.2e}", "Nao atendida; ha indicio de variancia nao constante."],
        ["Multicolinearidade", f"Maior VIF = {format_num(max(results['VIF_top'].values()), 2)}", "Sem multicolinearidade severa; VIF abaixo de 5 nas principais variaveis."],
    ]
    add_table(doc, ["Suposicao", "Evidencia", "Conclusao"], diag_rows, widths=[2500, 3000, 3860])

    doc.add_heading("Recomendacoes sobre as variaveis do modelo", level=1)
    doc.add_heading("Pontos positivos", level=2)
    add_bullet(doc, "A base traz variaveis cadastrais, patrimoniais e de renda que tem relacao estatistica relevante com o score.")
    add_bullet(doc, "Valor do imovel, salario, tempo no ultimo servico e posse de ativos melhoram a capacidade explicativa do modelo.")
    add_bullet(doc, "A ausencia de nulos facilita a implantacao inicial e reduz a necessidade de imputacao.")
    doc.add_heading("Pontos negativos e limites", level=2)
    add_bullet(doc, "As variaveis nao sao suficientes para uma decisao final de credito sem outras fontes: faltam historico de pagamento, atrasos, renda comprometida, dividas existentes, utilizacao de limite, consultas recentes e dados de bureau.")
    add_bullet(doc, "Algumas relacoes podem refletir proxies socioeconomicos; recomenda-se auditoria de vies, especialmente em variaveis como regiao de moradia, sexo e estado civil.")
    add_bullet(doc, "Como normalidade e homocedasticidade nao foram atendidas, inferencias de significancia devem ser usadas com cautela; para previsao, as metricas de teste e validacao cruzada sao mais relevantes.")

    doc.add_heading("Conclusao na comparacao com a arvore", level=1)
    add_number(doc, "A regressao linear e recomendada como modelo base por entregar menor erro e maior explicabilidade.")
    add_number(doc, "A arvore e inferior neste teste, mas ajuda a validar se ha regras nao lineares simples; neste caso, nao houve ganho.")
    add_number(doc, "Alternativas recomendadas para evolucao: Random Forest, Gradient Boosting, XGBoost/LightGBM, regressao robusta, regressao quantil e calibracao por faixas de risco.")
    add_number(doc, "Antes de usar em producao, definir politica de aprovacao por faixas, monitorar estabilidade populacional, testar vies e acompanhar performance em safras futuras.")

    doc.add_heading("Simulador do modelo", level=1)
    doc.add_paragraph("O script entregue treina a regressao linear, salva o pipeline e permite simular novos clientes por linha de comando:")
    p = doc.add_paragraph()
    run = p.add_run("python modelo_credit_scoring_quantum.py --simulate")
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    doc.add_paragraph("A classificacao sugerida por faixa foi definida como regra operacional inicial: score abaixo de 400 = alto risco; 400 a 599 = risco medio; 600 ou mais = baixo risco. Os pontos de corte devem ser calibrados pela taxa real de inadimplencia e pelo apetite de risco da Quantum Finance.")

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Quantum Finance - Applied Statistics").font.size = Pt(9)

    OUT_DOCX.parent.mkdir(exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
